from __future__ import annotations

import re
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

try:
    from PIL import Image as PILImage
except Exception:  # pragma: no cover - Pillow is a runtime dependency, but keep import safe.
    PILImage = None


_PDF_FONTS_CACHE: Optional[Tuple[str, str]] = None


def _register_pdf_unicode_fonts() -> Tuple[str, str]:
    global _PDF_FONTS_CACHE
    if _PDF_FONTS_CACHE:
        return _PDF_FONTS_CACHE

    font_pairs = [
        (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")),
        (Path("C:/Windows/Fonts/calibri.ttf"), Path("C:/Windows/Fonts/calibrib.ttf")),
        (Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"), Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")),
        (Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"), Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf")),
    ]
    for regular_path, bold_path in font_pairs:
        if not regular_path.exists() or not bold_path.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("DICORUnicode", str(regular_path)))
            pdfmetrics.registerFont(TTFont("DICORUnicode-Bold", str(bold_path)))
            _PDF_FONTS_CACHE = ("DICORUnicode", "DICORUnicode-Bold")
            return _PDF_FONTS_CACHE
        except Exception:
            continue

    _PDF_FONTS_CACHE = ("Helvetica", "Helvetica-Bold")
    return _PDF_FONTS_CACHE


SECTION_DEFINITIONS: List[Dict[str, Any]] = [
    {"item": 1, "key": "painel", "title": "PAINEL"},
    {"item": 2, "key": "fotos_lideres", "title": "FOTOS DOS LÍDERES"},
    {"item": 3, "key": "fotos_membros", "title": "FOTOS DOS MEMBROS"},
    {"item": 4, "key": "radio", "title": "RÁDIO"},
    {"item": 5, "key": "localizacao", "title": "LOCALIZAÇÃO", "staged": True},
    {"item": 6, "key": "crimes", "title": "CRIMES DA COMUNIDADE"},
    {"item": 7, "key": "bau_lider", "title": "BAÚ DE LÍDER"},
    {"item": 8, "key": "bau_membros", "title": "BAÚ DE MEMBROS"},
    {"item": 9, "key": "rota_farm", "title": "ROTA DE FARM"},
    {"item": 10, "key": "rota_producao", "title": "ROTA DE PRODUÇÃO"},
    {"item": 11, "key": "ingredientes_produtos", "title": "INGREDIENTES E PRODUTOS", "staged": True},
    {"item": 12, "key": "informante", "title": "INFORMANTE", "staged": True},
    {"item": 13, "key": "residencia_lider", "title": "RESIDÊNCIA DO LÍDER"},
]

TWO_STAGE_ITEMS = {5, 11, 12}

SECTION_KEY_ALIASES: Dict[str, int] = {
    "painel": 1,
    "fotos lideres": 2,
    "fotos dos lideres": 2,
    "fotos lider": 2,
    "fotos do lider": 2,
    "liderancas": 2,
    "lideranca": 2,
    "fotos membros": 3,
    "fotos dos membros": 3,
    "integrantes": 3,
    "membros": 3,
    "radio": 4,
    "localizacao": 5,
    "localização": 5,
    "crimes": 6,
    "crimes comunidade": 6,
    "bau lider": 7,
    "bau de lider": 7,
    "baus lider": 7,
    "baus_lider": 7,
    "bau membros": 8,
    "bau de membros": 8,
    "baus membros": 8,
    "baus_membros": 8,
    "rota farm": 9,
    "farm": 9,
    "rota_farm": 9,
    "rota producao": 10,
    "rota produção": 10,
    "producao": 10,
    "produção": 10,
    "rota_producao": 10,
    "ingredientes": 11,
    "produtos": 11,
    "materiais": 11,
    "ingredientes produtos": 11,
    "ingredientes_produtos": 11,
    "informante": 12,
    "informantes": 12,
    "residencia": 13,
    "residência": 13,
    "residencia lider": 13,
    "residência líder": 13,
    "residencia_lider": 13,
}

ADMIN_MESSAGE_PATTERNS = (
    "tarefa guiada",
    "tarefas da investigacao",
    "tarefas da investigação",
    "gerenciamento de tarefas",
    "checklist atual",
    "finalizar tarefa",
    "concluir tarefa",
    "material minimo completo",
    "material mínimo completo",
    "envie o material neste topico",
    "envie o material neste tópico",
    "aguarda conclusao",
    "aguarda conclusão",
    "concluido — item",
    "concluído — item",
    "etapa 1/2 liberada",
    "reanalisei",
    "reanalisar mesa",
    "gerar dossi",
    "[dicor] etapa",
    "encerramento confirmado",
)


def normalize_text(value: Any) -> str:
    text = str(value if value is not None else "").strip().lower()
    repl = str.maketrans("áàâãäéèêëíìîïóòôõöúùûüçñ", "aaaaaeeeeiiiiooooouuuucn")
    text = text.translate(repl)
    text = re.sub(r"[_\-–—•|]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def safe_text(value: Any, default: str = "Não informado") -> str:
    text = str(value if value is not None else "").replace("\r", "").strip()
    return text or default


def section_by_item(item: int) -> Dict[str, Any]:
    for sec in SECTION_DEFINITIONS:
        if int(sec["item"]) == int(item):
            return sec
    raise KeyError(f"Item inválido: {item}")


def classify_section(value: Any) -> int:
    normalized = normalize_text(value)
    if not normalized:
        return 0
    match = re.search(r"\b(?:item\s*)?([1-9]|1[0-3])\b", normalized)
    if match:
        try:
            item = int(match.group(1))
            if 1 <= item <= 13:
                return item
        except Exception:
            pass
    for alias, item in sorted(SECTION_KEY_ALIASES.items(), key=lambda kv: len(kv[0]), reverse=True):
        if alias in normalized:
            return item
    return 0


def is_administrative_message(message: Dict[str, Any]) -> bool:
    author_bot = bool(message.get("author_bot") or message.get("bot"))
    text = normalize_text(message.get("conteudo") or message.get("content") or message.get("texto") or "")
    if not text:
        return author_bot
    if any(pattern in text for pattern in ADMIN_MESSAGE_PATTERNS):
        return True
    return bool(author_bot and ("dicor" in text or "tarefa" in text or "painel" in text))


def _stage_for_message(message: Dict[str, Any], item: int, state: Dict[str, Any]) -> int:
    if item in TWO_STAGE_ITEMS:
        for field in ("etapa", "stage", "fase"):
            raw_stage = message.get(field)
            if raw_stage is None:
                continue
            match = re.search(r"\b([12])\b", normalize_text(raw_stage))
            if match:
                return int(match.group(1))
        text = normalize_text(
            " ".join(
                str(message.get(field) or "")
                for field in (
                    "conteudo",
                    "content",
                    "texto",
                    "origem",
                    "topico",
                    "topic",
                    "arquivo",
                    "filename",
                    "url",
                    "mensagem_url",
                    "jump_url",
                )
            )
        )
        if "etapa 2" in text or "2/2" in text:
            return 2
        if "etapa 1" in text or "1/2" in text:
            return 1
    etapas = state.get("etapas_por_item") if isinstance(state.get("etapas_por_item"), dict) else {}
    try:
        current = int(etapas.get(str(item)) or 1)
    except Exception:
        current = 1
    return 1 if item not in TWO_STAGE_ITEMS else min(max(current, 1), 2)


def _clean_message(message: Dict[str, Any], item: int, state: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": message.get("id") or message.get("mensagem_id") or "",
        "author": safe_text(message.get("autor") or message.get("author"), "Autor não identificado"),
        "author_id": message.get("autor_id") or message.get("author_id") or "",
        "date": safe_text(message.get("data") or message.get("created_at"), ""),
        "origin": safe_text(message.get("origem") or message.get("topico") or message.get("topic"), ""),
        "url": str(message.get("url") or message.get("mensagem_url") or message.get("jump_url") or "").strip(),
        "content": safe_text(message.get("conteudo") or message.get("content") or message.get("texto"), ""),
        "stage": _stage_for_message(message, item, state),
    }


def _evidence_type(evidence: Dict[str, Any]) -> str:
    tipo = normalize_text(evidence.get("tipo") or evidence.get("content_type") or evidence.get("arquivo") or "")
    if any(x in tipo for x in ("image", "imagem", "png", "jpg", "jpeg", "webp", "gif")):
        return "imagem"
    if any(x in tipo for x in ("video", "mp4", "mov", "webm", "avi")):
        return "video"
    if any(x in tipo for x in ("pdf", "doc", "document", "texto")):
        return "arquivo"
    return "link" if str(evidence.get("url") or "").startswith(("http://", "https://")) else "arquivo"


def _clean_evidence(evidence: Dict[str, Any], number: int, item: int, state: Dict[str, Any]) -> Dict[str, Any]:
    local = str(evidence.get("local") or evidence.get("path") or "").strip()
    return {
        "number": number,
        "id": f"E{number:03d}",
        "item": item,
        "section": section_by_item(item)["key"],
        "stage": _stage_for_message(evidence, item, state),
        "type": _evidence_type(evidence),
        "filename": safe_text(evidence.get("arquivo") or evidence.get("filename"), "arquivo"),
        "url": str(evidence.get("url") or evidence.get("proxy_url") or "").strip(),
        "local": local,
        "message_id": evidence.get("mensagem_id") or evidence.get("id") or "",
        "message_url": str(evidence.get("mensagem_url") or evidence.get("jump_url") or "").strip(),
        "author": safe_text(evidence.get("autor") or evidence.get("author"), "Autor não identificado"),
        "date": safe_text(evidence.get("data") or evidence.get("created_at"), ""),
        "origin": safe_text(evidence.get("origem") or evidence.get("topico") or evidence.get("topic"), ""),
    }


def _dedupe_key(*values: Any) -> Tuple[str, ...]:
    return tuple(normalize_text(v) for v in values)


def build_operational_dossier_payload(
    dados: Dict[str, Any],
    mesa_state: Optional[Dict[str, Any]] = None,
    *,
    assets: Optional[Dict[str, str]] = None,
    signatures: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    state = dict(mesa_state or {})
    payload: Dict[str, Any] = {
        "version": "v123",
        "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "metadata": {
            "processo": safe_text(dados.get("processo")),
            "numero_investigacao": safe_text(dados.get("numero_investigacao")),
            "nome_operacao": safe_text(dados.get("nome_operacao") or dados.get("nome")),
            "comunidade": safe_text(dados.get("comunidade")),
            "faccao": safe_text(dados.get("faccao")),
            "canal_id": safe_text(dados.get("canal_id"), ""),
            "canal_nome": safe_text(dados.get("canal_nome"), ""),
            "guild_id": safe_text(dados.get("guild_id"), ""),
            "guild_nome": safe_text(dados.get("guild_nome"), ""),
            "data_abertura": safe_text(dados.get("data_abertura")),
            "data_encerramento": safe_text(dados.get("data_encerramento")),
            "agente_encerramento": safe_text(dados.get("agente_encerramento")),
            "mesa_criada_por": safe_text((dados.get("mesa") or {}).get("autor_nome") or dados.get("delegado_responsavel")),
            "reabrir_url": str(dados.get("reabrir_url") or "").strip(),
        },
        "task_state": state,
        "assets": dict(assets or {}),
        "signatures": list(signatures or []),
        "sections": [],
        "evidence_index": [],
        "people": {
            "liderancas": list(dados.get("liderancas") or []),
            "integrantes": list(dados.get("integrantes") or []),
            "informantes": list(dados.get("informantes") or []),
        },
        "statistics": dict(dados.get("estatisticas") or {}),
    }

    sections: Dict[int, Dict[str, Any]] = {}
    for sec in SECTION_DEFINITIONS:
        item = int(sec["item"])
        sections[item] = {
            "item": item,
            "key": sec["key"],
            "title": sec["title"],
            "staged": bool(sec.get("staged")),
            "current_stage": int((state.get("etapas_por_item") or {}).get(str(item), 1)) if isinstance(state.get("etapas_por_item"), dict) else 1,
            "completed": str(item) in (state.get("concluidos") or {}) if isinstance(state.get("concluidos"), dict) else False,
            "tasks": {},
            "messages": [],
            "evidence": [],
            "summary": "",
        }

    task_records = state.get("tarefas_guiadas") if isinstance(state.get("tarefas_guiadas"), dict) else {}
    for task_key, task in task_records.items():
        if not isinstance(task, dict):
            continue
        try:
            item = int(task.get("item") or str(task_key).split(":", 1)[0])
            stage = int(task.get("etapa") or str(task_key).split(":", 1)[1])
        except Exception:
            continue
        if item in sections:
            sections[item]["tasks"][str(stage)] = dict(task)

    seen_messages = set()
    for message in list(dados.get("mensagens") or []):
        if not isinstance(message, dict) or is_administrative_message(message):
            continue
        item = classify_section(message.get("item") or message.get("topico") or message.get("origem"))
        if item not in sections:
            continue
        cleaned = _clean_message(message, item, state)
        key = _dedupe_key(cleaned.get("id"), cleaned.get("url"), cleaned.get("content"))
        if key in seen_messages:
            continue
        seen_messages.add(key)
        sections[item]["messages"].append(cleaned)

    evidence_number = 1
    seen_evidence = set()
    for evidence in list(dados.get("evidencias") or []):
        if not isinstance(evidence, dict):
            continue
        item = classify_section(evidence.get("item") or evidence.get("topico") or evidence.get("origem"))
        if item not in sections:
            continue
        key = _dedupe_key(evidence.get("mensagem_id"), evidence.get("url") or evidence.get("proxy_url"), evidence.get("local"), evidence.get("arquivo"))
        if key in seen_evidence:
            continue
        seen_evidence.add(key)
        cleaned = _clean_evidence(evidence, evidence_number, item, state)
        sections[item]["evidence"].append(cleaned)
        payload["evidence_index"].append(cleaned)
        evidence_number += 1

    resumos = dados.get("resumos") if isinstance(dados.get("resumos"), dict) else {}
    summary_aliases = {
        1: ("painel",),
        2: ("liderancas", "fotos_lideres"),
        3: ("integrantes", "fotos_membros"),
        4: ("radio",),
        5: ("localizacao",),
        6: ("crimes",),
        7: ("baus_lider", "bau_lider", "baus"),
        8: ("baus_membros", "bau_membros", "baus"),
        9: ("rota_farm", "farm"),
        10: ("rota_producao", "producao"),
        11: ("ingredientes_produtos", "materiais", "producao"),
        12: ("informantes", "informante"),
        13: ("residencia", "residencia_lider"),
    }
    for item, sec in sections.items():
        candidates = [safe_text(resumos.get(k), "") for k in summary_aliases.get(item, ())]
        candidates.extend(m["content"] for m in sec["messages"][:8])
        summary = "\n".join(dict.fromkeys([c.strip() for c in candidates if c and c.strip()]))
        sec["summary"] = summary[:3500] if summary else "Sem registro textual útil nesta seção; verifique as evidências e referências preservadas."

    payload["sections"] = [sections[i] for i in range(1, 14)]
    stats = payload["statistics"]
    stats["mensagens_v123"] = sum(len(s["messages"]) for s in payload["sections"])
    stats["evidencias_v123"] = len(payload["evidence_index"])
    stats["imagens_v123"] = len([e for e in payload["evidence_index"] if e["type"] == "imagem"])
    stats["videos_v123"] = len([e for e in payload["evidence_index"] if e["type"] == "video"])
    return payload


def _draw_image_fit(c: canvas.Canvas, path: str, x: float, y: float, max_w: float, max_h: float) -> Tuple[float, float]:
    p = Path(str(path or ""))
    if not p.exists() or not p.is_file():
        return 0, 0
    width, height = max_w, max_h
    if PILImage is not None:
        try:
            with PILImage.open(p) as img:
                iw, ih = img.size
                if iw > 0 and ih > 0:
                    scale = min(max_w / iw, max_h / ih)
                    width, height = iw * scale, ih * scale
        except Exception:
            pass
    c.drawImage(str(p), x, y + (max_h - height), width=width, height=height, preserveAspectRatio=True, mask="auto")
    return width, height


def _wrap_lines(text: Any, width: int = 96) -> List[str]:
    output: List[str] = []
    for raw in safe_text(text, "").splitlines():
        raw = raw.strip()
        if not raw:
            output.append("")
            continue
        output.extend(textwrap.wrap(raw, width=width, break_long_words=False, replace_whitespace=False) or [""])
    return output or ["Não informado"]


def _pdf_page(c: canvas.Canvas, payload: Dict[str, Any], title: str, *, continuation: bool = False) -> Tuple[float, float, float, float]:
    width, height = A4
    regular_font, bold_font = _register_pdf_unicode_fonts()
    assets = payload.get("assets") or {}
    frame = assets.get("moldura")
    watermark = assets.get("marca_dagua")
    pf = assets.get("brasao_pf")
    dicor = assets.get("brasao_dicor")
    c.setFillColor(colors.HexColor("#6F716D"))
    c.rect(0, 0, width, height, stroke=0, fill=1)
    if frame and Path(str(frame)).exists():
        c.drawImage(str(frame), 0, 0, width=width, height=height, preserveAspectRatio=False, mask="auto")
    logo = 2.55 * cm
    if pf and Path(str(pf)).exists():
        c.drawImage(str(pf), 1.35 * cm, height - 3.45 * cm, width=logo, height=logo, preserveAspectRatio=True, mask="auto")
    if dicor and Path(str(dicor)).exists():
        c.drawImage(str(dicor), width - 1.35 * cm - logo, height - 3.45 * cm, width=logo, height=logo, preserveAspectRatio=True, mask="auto")
    c.setFillColor(colors.HexColor("#D7B66A"))
    c.setFont(bold_font, 22)
    c.drawCentredString(width / 2, height - 1.45 * cm, "POLÍCIA FEDERAL • DICOR")
    c.setFillColor(colors.white)
    c.setFont(bold_font, 12)
    c.drawCentredString(width / 2, height - 2.02 * cm, "DOSSIÊ OPERACIONAL DE INVESTIGAÇÃO")

    left = 1.34 * cm
    right = width - 1.34 * cm
    top = height - 4.2 * cm
    bottom = 1.35 * cm
    c.saveState()
    try:
        c.setFillAlpha(0.84)
    except Exception:
        pass
    c.setFillColor(colors.HexColor("#0C1826"))
    c.roundRect(left - 0.1 * cm, bottom, right - left + 0.2 * cm, top - bottom + 0.35 * cm, 7, fill=1, stroke=0)
    c.restoreState()
    if watermark and Path(str(watermark)).exists():
        c.saveState()
        try:
            c.setFillAlpha(0.18)
            c.setStrokeAlpha(0.18)
        except Exception:
            pass
        c.drawImage(str(watermark), (width - 10 * cm) / 2, (height - 12 * cm) / 2 - 0.8 * cm, width=10 * cm, height=12 * cm, preserveAspectRatio=True, mask="auto")
        c.restoreState()
    bar_y = top - 0.8 * cm
    c.setFillColor(colors.HexColor("#003B6F"))
    c.rect(left, bar_y, right - left, 0.7 * cm, fill=1, stroke=0)
    c.setStrokeColor(colors.HexColor("#D7B66A"))
    c.rect(left, bar_y, right - left, 0.7 * cm, fill=0, stroke=1)
    c.setFillColor(colors.white)
    c.setFont(bold_font, 11.2)
    header = title + (" — CONTINUAÇÃO" if continuation else "")
    c.drawString(left + 0.18 * cm, bar_y + 0.22 * cm, header[:88])
    c.setFont(regular_font, 7.3)
    c.drawRightString(right - 0.15 * cm, bottom - 0.32 * cm, f"V123 • Página {c.getPageNumber()}")
    return left, right, bar_y - 0.42 * cm, bottom + 0.18 * cm


def generate_pdf(payload: Dict[str, Any], output_path: Path) -> None:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(output_path), pagesize=A4)
    regular_font, bold_font = _register_pdf_unicode_fonts()
    left, right, y, bottom = _pdf_page(c, payload, "CAPA E IDENTIFICAÇÃO")
    meta = payload.get("metadata") or {}
    c.setFillColor(colors.HexColor("#D7B66A"))
    c.setFont(bold_font, 16)
    c.drawCentredString((left + right) / 2, y - 0.2 * cm, "DOSSIÊ OPERACIONAL")
    y -= 1.05 * cm
    rows = [
        ("Processo", meta.get("processo")),
        ("Investigação", meta.get("numero_investigacao")),
        ("Operação", meta.get("nome_operacao")),
        ("Comunidade", meta.get("comunidade")),
        ("Facção/Estrutura", meta.get("faccao")),
        ("Mesa", f"{meta.get('canal_nome')} ({meta.get('canal_id')})"),
        ("Mesa criada por", meta.get("mesa_criada_por")),
        ("Abertura", meta.get("data_abertura")),
        ("Encerramento", meta.get("data_encerramento")),
        ("Agente do encerramento", meta.get("agente_encerramento")),
    ]
    for label, value in rows:
        if y < bottom + 0.5 * cm:
            c.showPage()
            left, right, y, bottom = _pdf_page(c, payload, "CAPA E IDENTIFICAÇÃO", continuation=True)
        c.setFillColor(colors.HexColor("#D7B66A"))
        c.setFont(bold_font, 8.8)
        c.drawString(left + 0.25 * cm, y, f"{label}:")
        c.setFillColor(colors.white)
        c.setFont(regular_font, 8.8)
        c.drawString(left + 5.2 * cm, y, safe_text(value)[:110])
        y -= 0.42 * cm

    y -= 0.28 * cm
    c.setFillColor(colors.HexColor("#D7B66A"))
    c.setFont(bold_font, 10.5)
    c.drawString(left + 0.25 * cm, y, "ORDEM OFICIAL DAS SEÇÕES")
    y -= 0.45 * cm
    c.setFont(regular_font, 8.2)
    c.setFillColor(colors.white)
    for sec in payload.get("sections") or []:
        c.drawString(left + 0.45 * cm, y, f"{sec['item']:02d}. {sec['title']}")
        y -= 0.32 * cm

    for sec in payload.get("sections") or []:
        c.showPage()
        left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}")
        c.setFillColor(colors.HexColor("#D7B66A"))
        c.setFont(bold_font, 9)
        status = "CONCLUÍDA" if sec.get("completed") else "NÃO CONCLUÍDA"
        stage = f" • etapa atual {sec.get('current_stage', 1)}/2" if sec.get("staged") else ""
        c.drawString(left + 0.2 * cm, y, f"Status da tarefa: {status}{stage}")
        y -= 0.48 * cm

        c.setFillColor(colors.white)
        c.setFont(regular_font, 8.35)
        for line in _wrap_lines(sec.get("summary"), 98):
            if y < bottom + 0.45 * cm:
                c.showPage()
                left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}", continuation=True)
            c.drawString(left + 0.25 * cm, y, line[:140])
            y -= 0.34 * cm

        messages = list(sec.get("messages") or [])
        if messages:
            if y < bottom + 1.2 * cm:
                c.showPage()
                left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}", continuation=True)
            c.setFillColor(colors.HexColor("#D7B66A"))
            c.setFont(bold_font, 8.4)
            c.drawString(left + 0.25 * cm, y, "Mensagens úteis preservadas")
            y -= 0.38 * cm
            c.setFillColor(colors.white)
            c.setFont(regular_font, 7.2)
            for msg in messages[:12]:
                stage = f"Etapa {msg.get('stage', 1)}/2 • " if sec.get("staged") else ""
                msg_text = f"{stage}{msg.get('date')} • {msg.get('author')} • msg {msg.get('id')}: {msg.get('content')}"
                for line in _wrap_lines(msg_text, 116):
                    if y < bottom + 0.35 * cm:
                        c.showPage()
                        left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}", continuation=True)
                        c.setFillColor(colors.white)
                        c.setFont(regular_font, 7.2)
                    c.drawString(left + 0.35 * cm, y, line[:150])
                    y -= 0.3 * cm

        images = [ev for ev in sec.get("evidence") or [] if ev.get("type") == "imagem" and ev.get("local") and Path(str(ev.get("local"))).exists()]
        for ev in images[:10]:
            if y < bottom + 5.2 * cm:
                c.showPage()
                left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}", continuation=True)
            y -= 0.15 * cm
            _draw_image_fit(c, str(ev.get("local")), left + 0.3 * cm, y - 4.2 * cm, right - left - 0.6 * cm, 4.0 * cm)
            y -= 4.45 * cm
            c.setFillColor(colors.HexColor("#D7B66A"))
            c.setFont(regular_font, 7.1)
            stage = f" • etapa {ev.get('stage', 1)}" if sec.get("staged") else ""
            c.drawString(left + 0.3 * cm, y, f"{ev['id']} • {ev.get('filename')}{stage} • tópico: {ev.get('origin')} • msg: {ev.get('message_id')}")
            y -= 0.32 * cm

        non_images = [ev for ev in sec.get("evidence") or [] if ev.get("type") != "imagem"]
        if non_images:
            if y < bottom + 1.2 * cm:
                c.showPage()
                left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}", continuation=True)
            c.setFillColor(colors.HexColor("#D7B66A"))
            c.setFont(bold_font, 8.4)
            c.drawString(left + 0.25 * cm, y, "Arquivos, vídeos e links preservados")
            y -= 0.38 * cm
            c.setFillColor(colors.white)
            c.setFont(regular_font, 7.2)
            for ev in non_images[:20]:
                if y < bottom + 0.35 * cm:
                    c.showPage()
                    left, right, y, bottom = _pdf_page(c, payload, f"{sec['item']:02d}. {sec['title']}", continuation=True)
                stage = f" • etapa {ev.get('stage', 1)}" if sec.get("staged") else ""
                c.drawString(left + 0.35 * cm, y, f"{ev['id']} • {ev.get('type')}{stage} • {ev.get('filename')} • msg: {ev.get('message_id')} • {ev.get('url')[:80]}")
                y -= 0.3 * cm

    c.showPage()
    left, right, y, bottom = _pdf_page(c, payload, "ÍNDICE DE EVIDÊNCIAS E ASSINATURAS")
    c.setFillColor(colors.white)
    c.setFont(regular_font, 7.1)
    for ev in payload.get("evidence_index") or []:
        if y < bottom + 0.3 * cm:
            c.showPage()
            left, right, y, bottom = _pdf_page(c, payload, "ÍNDICE DE EVIDÊNCIAS E ASSINATURAS", continuation=True)
            c.setFillColor(colors.white)
            c.setFont(regular_font, 7.1)
        c.drawString(left + 0.25 * cm, y, f"{ev['id']} • Item {ev['item']:02d} • {ev.get('filename')} • {ev.get('author')} • {ev.get('date')} • msg {ev.get('message_id')}")
        y -= 0.28 * cm

    y -= 0.35 * cm
    c.setFillColor(colors.HexColor("#D7B66A"))
    c.setFont(bold_font, 9)
    c.drawString(left + 0.25 * cm, y, "Assinaturas institucionais")
    y -= 0.42 * cm
    for sig in payload.get("signatures") or []:
        c.setFillColor(colors.white)
        c.setFont(bold_font, 8.3)
        c.drawString(left + 0.45 * cm, y, f"{safe_text(sig.get('titulo'), 'Assinatura')} — {safe_text(sig.get('nome'))}")
        img = sig.get("imagem") or sig.get("arquivo")
        if img and Path(str(img)).exists():
            _draw_image_fit(c, str(img), right - 5.0 * cm, y - 1.0 * cm, 4.2 * cm, 1.0 * cm)
        y -= 1.15 * cm

    c.save()


def _docx_add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 43, 91)


def _docx_add_picture_if_exists(paragraph, path: Any, *, width: float) -> bool:
    p = Path(str(path or ""))
    if not p.exists() or not p.is_file():
        return False
    try:
        paragraph.add_run().add_picture(str(p), width=Inches(width))
        return True
    except Exception:
        return False


def _docx_fit_picture_inches(path: Any, *, max_width: float, max_height: float) -> Tuple[float, float]:
    p = Path(str(path or ""))
    if PILImage is None or not p.exists() or not p.is_file():
        return max_width, max_height
    try:
        with PILImage.open(p) as img:
            iw, ih = img.size
            if iw <= 0 or ih <= 0:
                return max_width, max_height
            scale = min(max_width / float(iw), max_height / float(ih))
            return max(0.1, iw * scale), max(0.1, ih * scale)
    except Exception:
        return max_width, max_height


def _docx_anchor_picture_behind_text(shape: Any, *, x_inches: float, y_inches: float) -> bool:
    try:
        inline = shape._inline
        inline.tag = qn("wp:anchor")
        for attr, value in {
            "distT": "0",
            "distB": "0",
            "distL": "0",
            "distR": "0",
            "simplePos": "0",
            "relativeHeight": "0",
            "behindDoc": "1",
            "locked": "0",
            "layoutInCell": "1",
            "allowOverlap": "1",
        }.items():
            inline.set(attr, value)

        removable = {
            qn("wp:simplePos"),
            qn("wp:positionH"),
            qn("wp:positionV"),
            qn("wp:wrapNone"),
            qn("wp:wrapSquare"),
            qn("wp:wrapTight"),
            qn("wp:wrapThrough"),
            qn("wp:wrapTopAndBottom"),
        }
        for child in list(inline):
            if child.tag in removable:
                inline.remove(child)

        simple_pos = OxmlElement("wp:simplePos")
        simple_pos.set("x", "0")
        simple_pos.set("y", "0")

        position_h = OxmlElement("wp:positionH")
        position_h.set("relativeFrom", "page")
        offset_h = OxmlElement("wp:posOffset")
        offset_h.text = str(int(x_inches * 914400))
        position_h.append(offset_h)

        position_v = OxmlElement("wp:positionV")
        position_v.set("relativeFrom", "page")
        offset_v = OxmlElement("wp:posOffset")
        offset_v.text = str(int(y_inches * 914400))
        position_v.append(offset_v)

        wrap_none = OxmlElement("wp:wrapNone")
        inline.insert(0, simple_pos)
        inline.insert(1, position_h)
        inline.insert(2, position_v)
        doc_pr_idx = next((idx for idx, child in enumerate(inline) if child.tag == qn("wp:docPr")), len(inline))
        inline.insert(doc_pr_idx, wrap_none)
        return True
    except Exception:
        return False


def _docx_add_page_background_assets(doc: Document, payload: Dict[str, Any]) -> None:
    assets = payload.get("assets") or {}
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    frame = assets.get("moldura")
    if frame and Path(str(frame)).exists():
        try:
            shape = paragraph.add_run().add_picture(str(frame), width=Inches(8.27), height=Inches(11.69))
            _docx_anchor_picture_behind_text(shape, x_inches=0.0, y_inches=0.0)
        except Exception:
            pass

    watermark = assets.get("marca_dagua")
    if watermark and Path(str(watermark)).exists():
        try:
            shape = paragraph.add_run().add_picture(str(watermark), width=Inches(3.7), height=Inches(4.4))
            _docx_anchor_picture_behind_text(shape, x_inches=2.28, y_inches=3.62)
        except Exception:
            pass


def _docx_add_institutional_band(doc: Document, payload: Dict[str, Any], title: str) -> None:
    assets = payload.get("assets") or {}
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(1.15), Inches(4.9), Inches(1.15))
    for idx, width in enumerate(widths):
        cell = table.rows[0].cells[idx]
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left_p = table.rows[0].cells[0].paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _docx_add_picture_if_exists(left_p, assets.get("brasao_pf"), width=0.82)
    center_p = table.rows[0].cells[1].paragraphs[0]
    center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = center_p.add_run("POLÍCIA FEDERAL • DICOR\n")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 43, 91)
    sub = center_p.add_run(title)
    sub.bold = True
    sub.font.name = "Arial"
    sub.font.size = Pt(9)
    sub.font.color.rgb = RGBColor(138, 101, 0)
    right_p = table.rows[0].cells[2].paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _docx_add_picture_if_exists(right_p, assets.get("brasao_dicor"), width=0.82)

    watermark = assets.get("marca_dagua")
    if watermark and Path(str(watermark)).exists():
        wm_p = doc.add_paragraph()
        wm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_add_picture_if_exists(wm_p, watermark, width=1.45)


def _docx_add_paragraph(doc: Document, text: Any) -> None:
    for line in _wrap_lines(text, 110):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def _docx_add_table(doc: Document, rows: Iterable[Tuple[Any, Any]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Campo"
    table.rows[0].cells[1].text = "Informação"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = safe_text(label, "")
        cells[1].text = safe_text(value)


def generate_docx(payload: Dict[str, Any], output_path: Path) -> None:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    _docx_add_page_background_assets(doc, payload)

    _docx_add_institutional_band(doc, payload, "DOSSIÊ OPERACIONAL DE INVESTIGAÇÃO")
    title = doc.add_heading("DOSSIÊ OPERACIONAL DE INVESTIGAÇÃO", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = payload.get("metadata") or {}
    _docx_add_table(doc, [
        ("Processo", meta.get("processo")),
        ("Investigação", meta.get("numero_investigacao")),
        ("Operação", meta.get("nome_operacao")),
        ("Comunidade", meta.get("comunidade")),
        ("Facção/Estrutura", meta.get("faccao")),
        ("Mesa", f"{meta.get('canal_nome')} ({meta.get('canal_id')})"),
        ("Mesa criada por", meta.get("mesa_criada_por")),
        ("Abertura", meta.get("data_abertura")),
        ("Encerramento", meta.get("data_encerramento")),
        ("Agente do encerramento", meta.get("agente_encerramento")),
    ])

    _docx_add_heading(doc, "Ordem oficial das seções", 1)
    for sec_payload in payload.get("sections") or []:
        _docx_add_paragraph(doc, f"{sec_payload['item']:02d}. {sec_payload['title']}")

    for sec_payload in payload.get("sections") or []:
        doc.add_section(WD_SECTION.NEW_PAGE)
        _docx_add_institutional_band(doc, payload, f"{sec_payload['item']:02d}. {sec_payload['title']}")
        _docx_add_heading(doc, f"{sec_payload['item']:02d}. {sec_payload['title']}", 1)
        stage = f" • etapa atual {sec_payload.get('current_stage', 1)}/2" if sec_payload.get("staged") else ""
        _docx_add_paragraph(doc, f"Status da tarefa: {'CONCLUÍDA' if sec_payload.get('completed') else 'NÃO CONCLUÍDA'}{stage}")
        _docx_add_paragraph(doc, sec_payload.get("summary"))

        if sec_payload.get("messages"):
            _docx_add_heading(doc, "Mensagens úteis preservadas", 2)
            for msg in sec_payload.get("messages", [])[:12]:
                _docx_add_paragraph(doc, f"• {msg.get('date')} • {msg.get('author')} • msg {msg.get('id')}: {msg.get('content')}")

        if sec_payload.get("evidence"):
            _docx_add_heading(doc, "Evidências rastreáveis", 2)
            for ev in sec_payload.get("evidence", [])[:30]:
                _docx_add_paragraph(doc, f"{ev['id']} • {ev.get('type')} • {ev.get('filename')} • tópico: {ev.get('origin')} • autor: {ev.get('author')} • data: {ev.get('date')} • msg: {ev.get('message_id')}")
                if ev.get("type") == "imagem" and ev.get("local") and Path(str(ev.get("local"))).exists():
                    try:
                        img_width, img_height = _docx_fit_picture_inches(ev.get("local"), max_width=5.9, max_height=4.2)
                        doc.add_picture(str(ev.get("local")), width=Inches(img_width), height=Inches(img_height))
                    except Exception:
                        _docx_add_paragraph(doc, f"Imagem preservada em: {ev.get('local')}")

    doc.add_section(WD_SECTION.NEW_PAGE)
    _docx_add_institutional_band(doc, payload, "ÍNDICE DE EVIDÊNCIAS")
    _docx_add_heading(doc, "Índice de evidências", 1)
    for ev in payload.get("evidence_index") or []:
        _docx_add_paragraph(doc, f"{ev['id']} • Item {ev['item']:02d} • {ev.get('filename')} • {ev.get('message_url') or ev.get('url')}")

    if payload.get("signatures"):
        _docx_add_heading(doc, "Assinaturas institucionais", 1)
        for sig in payload.get("signatures") or []:
            _docx_add_paragraph(doc, f"{safe_text(sig.get('titulo'), 'Assinatura')} — {safe_text(sig.get('nome'))}")
            img = sig.get("imagem") or sig.get("arquivo")
            if img and Path(str(img)).exists():
                try:
                    doc.add_picture(str(img), width=Inches(2.6))
                except Exception:
                    pass

    doc.save(str(output_path))
